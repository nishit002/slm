"""
COMPLETE AI CURRICULUM GENERATOR - FULLY FIXED VERSION
========================================================
"""

import streamlit as st
import requests
import json
import time
import os
import re
from datetime import datetime
from io import BytesIO
from PIL import Image as PilImage

# PDF imports
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    st.warning("⚠️ PyPDF2 not installed - PDF upload disabled")

# ReportLab imports for PDF generation
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        Table, TableStyle, Image, KeepTogether
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    st.error("❌ ReportLab not installed - PDF generation disabled")

# DOCX imports for editable output
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx not installed - DOCX generation disabled")

# ============================================================================
# CONFIGURATION
# ============================================================================

# Grok API Configuration
DEFAULT_API_KEY = "xai-6QJwG3u6540lVZyXbFBArvLQ43ZyJsrnq65pyCWhxh5zXqNvtwe6LdTURbTwvE2sA3Uxlb9gn82Vamgu"
API_URL = "https://api.x.ai/v1/chat/completions"

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def initialize_session_state():
    """Initialize all session state variables"""
    defaults = {
        # API
        'api_key': DEFAULT_API_KEY,
        
        # Navigation
        'step': 'configuration',
        
        # Course Info
        'course_title': '',
        'course_code': '',
        'credits': 3,
        'target_audience': '',
        'learning_objectives': '',
        'prerequisites': '',
        'assessment_methods': '',
        
        # Program/Course Outcomes
        'program_objectives': '',
        'program_outcomes': '',
        'course_outcomes': '',
        'specialized_outcomes': '',

        # Document Customization
        'document_heading': '',
        'logo': None,
        
        # PDF Upload
        'uploaded_pdf_text': '',
        'pdf_processed': False,
        'extracted_structure': None,
        
        # Outline
        'raw_outline': '',
        'approved_outline': [],
        'outline_generated': False,
        
        # Content Generation
        'content': {},
        'content_status': {},
        'images': {},
        'image_prompts': {},
        'sections_to_process': [],
        'paused': False,
        'generation_start_time': None,
        
        # Compilation
        'compiled_files': {},
        'compile_type': 'Both (Separate + Complete)',
        'output_format': 'PDF',
        
        # UI Choices
        'upload_choice': 'Upload Syllabus PDF',
        'num_units': 4,
        'sections_per_unit': 8,
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

# ============================================================================
# Phase 1 Complete
# ============================================================================
print("Phase 1: Imports and Configuration loaded successfully")
"""
"""

# ============================================================================
# API HELPER FUNCTIONS
# ============================================================================

def get_api_headers():
    """Get API headers with current API key"""
    return {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.session_state.api_key}"
    }

def make_api_call(messages, retries=3, timeout=120, max_tokens=2000):
    """
    Make API call to Grok with detailed logging and error handling
    
    Args:
        messages: List of message dicts with role and content
        retries: Number of retry attempts
        timeout: Request timeout in seconds
        max_tokens: Maximum tokens in response
        
    Returns:
        str: API response content or None if failed
    """
    headers = get_api_headers()
    
    payload = {
        "messages": messages,
        "model": "grok-2-1212",
        "stream": False,
        "temperature": 0.3,
        "max_tokens": max_tokens
    }
    
    for attempt in range(retries):
        try:
            st.write(f"🔄 API Call Attempt {attempt + 1}/{retries}")
            st.write(f"📤 Sending to: {API_URL}")
            st.write(f"🎯 Model: grok-2-1212")
            st.write(f"📊 Max tokens: {max_tokens}")
            
            # Show first 200 chars of prompt
            if messages:
                prompt_preview = messages[-1].get('content', '')[:200]
                st.write(f"📝 Prompt preview: {prompt_preview}...")
            
            response = requests.post(API_URL, headers=headers, json=payload, timeout=timeout)
            
            st.write(f"📡 Response Status: {response.status_code}")
            
            # Show response time
            if response.headers:
                st.write(f"⏱️ Response time: {response.elapsed.total_seconds():.2f}s")
            
            response.raise_for_status()
            result = response.json()
            
            # Detailed response analysis
            if 'choices' in result and len(result['choices']) > 0:
                content = result['choices'][0]['message']['content']
                
                # Analyze response quality
                word_count = len(content.split())
                char_count = len(content)
                has_structure = any(keyword in content.upper() for keyword in ['INTRODUCTION', 'LEARNING OBJECTIVES', 'CHECK YOUR PROGRESS'])
                has_blooms = any(keyword in content.upper() for keyword in ['REMEMBER', 'UNDERSTAND', 'APPLY', 'ANALYZE', 'EVALUATE', 'CREATE'])
                
                st.write("✅ **API Response Analysis:**")
                st.write(f"   📊 Words: {word_count:,}")
                st.write(f"   📝 Characters: {char_count:,}")
                st.write(f"   📚 Has Structure: {'✅' if has_structure else '❌'}")
                st.write(f"   🎯 Has Bloom's Taxonomy: {'✅' if has_blooms else '❌'}")
                
                if word_count < 800:
                    st.warning(f"⚠️ Response seems short ({word_count} words). Expected 1,000+")
                    st.write("🔍 First 500 characters of response:")
                    st.code(content[:500])
                elif word_count < 1200:
                    st.warning(f"⚠️ Response shorter than expected ({word_count} words). Expected 1,000-1,500")
                else:
                    st.success(f"✅ Good response length: {word_count:,} words")
                
                return content
            else:
                st.error(f"❌ Unexpected response format")
                st.json(result)
                
        except requests.exceptions.HTTPError as e:
            st.error(f"❌ HTTP Error {e.response.status_code}")
            try:
                error_detail = e.response.json()
                st.json(error_detail)
            except:
                st.error(e.response.text[:500])
            
            if e.response.status_code == 401:
                st.error("🔑 Invalid API Key - Check your configuration")
                return None
            elif e.response.status_code == 429:
                st.warning("⏳ Rate limited - waiting before retry...")
                time.sleep(10)
            else:
                if attempt < retries - 1:
                    st.warning(f"⏳ Retrying in 3 seconds...")
                    time.sleep(3)
                    
        except requests.exceptions.Timeout:
            st.error(f"⏱️ Request timeout after {timeout}s")
            if attempt < retries - 1:
                st.warning("⏳ Retrying with longer timeout...")
                timeout += 30
                time.sleep(2)
                
        except requests.exceptions.ConnectionError as e:
            st.error(f"🔌 Connection error: {str(e)}")
            if attempt < retries - 1:
                st.warning("⏳ Retrying connection...")
                time.sleep(3)
                
        except Exception as e:
            st.error(f"❌ Unexpected error: {type(e).__name__}: {str(e)}")
            if attempt < retries - 1:
                time.sleep(2)
    
    st.error("❌ All API attempts failed - check logs above for details")
    return None

# ============================================================================
# PDF AND SYLLABUS PROCESSING
# ============================================================================

def extract_pdf_text(pdf_file):
    """
    Extract text from uploaded PDF file
    
    Args:
        pdf_file: Uploaded PDF file object
        
    Returns:
        str: Extracted text or None if failed
    """
    if not PYPDF2_AVAILABLE:
        st.error("PyPDF2 not installed - cannot extract PDF")
        return None
        
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
            
        st.session_state.uploaded_pdf_text = text
        st.session_state.pdf_processed = True
        st.success(f"✅ Extracted {len(text)} characters from {len(pdf_reader.pages)} pages")
        return text
        
    except Exception as e:
        st.error(f"❌ Error extracting PDF: {str(e)}")
        return None

def parse_syllabus_structure(text):
    """
    Parse syllabus text to extract course structure
    
    Args:
        text: Syllabus text content
        
    Returns:
        dict: Structured syllabus data with course_info and units
    """
    structure = {'course_info': {}, 'units': []}
    
    # Extract course information
    patterns = {
        'title': r'(?:Course|Subject)\s*(?:Title|Name)?\s*:?\s*(.+)',
        'code': r'(?:Course|Subject)\s*Code\s*:?\s*([A-Z0-9]+)',
        'credits': r'Credits?\s*:?\s*(\d+)',
    }
    
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            structure['course_info'][key] = match.group(1).strip()
    
    # Extract units with their topics
    unit_pattern = r'UNIT[\s-]*(\d+)\s*:?\s*(.+?)(?=UNIT[\s-]*\d+|$)'
    units = re.finditer(unit_pattern, text, re.IGNORECASE | re.DOTALL)
    
    for unit_match in units:
        unit_num = unit_match.group(1)
        unit_content = unit_match.group(2)
        
        # Extract unit title (first line)
        title_match = re.search(r'^(.+?)(?:\n|$)', unit_content)
        unit_title = title_match.group(1).strip() if title_match else f"Unit {unit_num}"
        
        # Extract topics (numbered lines)
        topics = []
        lines = unit_content.split('\n')
        for line in lines:
            if re.match(r'^\s*[\d.]+\s+(.+?)$', line):
                topic = re.match(r'^\s*[\d.]+\s+(.+?)$', line).group(1).strip()
                if 5 < len(topic) < 200:  # Filter reasonable topic lengths
                    topics.append(topic)
        
        structure['units'].append({
            'unit_number': int(unit_num),
            'unit_title': unit_title,
            'topics': topics
        })
    
    return structure

# ============================================================================
# TEXT CLEANING AND FORMATTING (FIXED LATEX HANDLING)
# ============================================================================

def clean_text_for_pdf(text):
    """
    Clean text for PDF with HTML tags for bold/italic and handle LaTeX equations
    
    Args:
        text: Raw text with markdown and LaTeX
        
    Returns:
        str: Cleaned text with HTML formatting
    """
    if not text:
        return ""
    
    # Convert LaTeX equations to readable format
    # Inline math: $$ ... $$ or $ ... $
    text = re.sub(r'\\$$(.*?)\\$$', r'[\1]', text)
    text = re.sub(r'\$([^\$]+?)\$', r'[\1]', text)
    
    # Display math: \[ ... \] or $$ ... $$
    text = re.sub(r'\\\[(.*?)\\\]', r'[\1]', text, flags=re.DOTALL)
    text = re.sub(r'\$\$(.*?)\$\$', r'[\1]', text, flags=re.DOTALL)
    
    # Common LaTeX commands to Unicode symbols
    latex_replacements = {
        r'\\leq': '≤',
        r'\\geq': '≥',
        r'\\neq': '≠',
        r'\\approx': '≈',
        r'\\equiv': '≡',
        r'\\times': '×',
        r'\\div': '÷',
        r'\\pm': '±',
        r'\\mp': '∓',
        r'\\sum': 'Σ',
        r'\\prod': 'Π',
        r'\\int': '∫',
        r'\\infty': '∞',
        r'\\partial': '∂',
        r'\\nabla': '∇',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\gamma': 'γ',
        r'\\delta': 'δ',
        r'\\epsilon': 'ε',
        r'\\theta': 'θ',
        r'\\lambda': 'λ',
        r'\\mu': 'μ',
        r'\\pi': 'π',
        r'\\sigma': 'σ',
        r'\\tau': 'τ',
        r'\\phi': 'φ',
        r'\\omega': 'ω',
        r'\\Gamma': 'Γ',
        r'\\Delta': 'Δ',
        r'\\Theta': 'Θ',
        r'\\Lambda': 'Λ',
        r'\\Sigma': 'Σ',
        r'\\Phi': 'Φ',
        r'\\Omega': 'Ω',
        r'\\rightarrow': '→',
        r'\\leftarrow': '←',
        r'\\Rightarrow': '⇒',
        r'\\Leftarrow': '⇐',
        r'\\leftrightarrow': '↔',
        r'\\Leftrightarrow': '⇔',
        r'\\forall': '∀',
        r'\\exists': '∃',
        r'\\in': '∈',
        r'\\notin': '∉',
        r'\\subset': '⊂',
        r'\\subseteq': '⊆',
        r'\\supset': '⊃',
        r'\\supseteq': '⊇',
        r'\\cup': '∪',
        r'\\cap': '∩',
        r'\\emptyset': '∅',
        r'\\_': '_',
        r'\\{': '{',
        r'\\}': '}',
        r'\\%': '%',
    }
    
    for latex, symbol in latex_replacements.items():
        text = re.sub(latex, symbol, text)
    
    # Handle subscripts and superscripts (simplified)
    text = re.sub(r'_\{([^}]+)\}', r'_\1', text)
    text = re.sub(r'\^\{([^}]+)\}', r'^\1', text)
    
    # Remove remaining LaTeX commands
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    
    # Handle markdown bold and italic
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    
    # Remove markdown headers
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    
    return text

# ============================================================================
# Phase 2 Complete
# ============================================================================
print("Phase 2: Helper Functions loaded successfully")
"""
"""

# ============================================================================
# CONTENT GENERATION
# ============================================================================

def generate_content(section_info, course_context):
    """
    Generate academic content for a section with image suggestions
    
    Args:
        section_info: Dict with section_number, section_title, description
        course_context: Dict with course details and outcomes
        
    Returns:
        str: Generated content
    """
    system_prompt = f"""You are an expert academic content developer for {course_context['target_audience']}.

Generate comprehensive content following eGyankosh standards:
- 4-5 pages (1,000-1,500 words)
- Grade 5 English, Academic tone
- Readability: 10-12
- Suggest 1-2 relevant image placements with descriptions (e.g., "Image: Diagram of organizational structure")

STRUCTURE (MUST INCLUDE ALL):
1. Introduction (2-3 paragraphs introducing the topic)
2. Learning Objectives (5-7 objectives mapped to Blooms Taxonomy: Remember, Understand, Apply, Analyze, Evaluate, Create)
3. Detailed Content (Main body with subsections, definitions, explanations)
4. Examples & Case Studies (Practical real-world examples)
5. CHECK YOUR PROGRESS (5-7 questions for self-assessment)
6. Summary (Concise recap of key points)
7. Key Terms (Glossary of important terms)

Map to:
- PO: {course_context.get('program_outcomes', 'N/A')}
- CO: {course_context.get('course_outcomes', 'N/A')}
- PSO: {course_context.get('specialized_outcomes', 'N/A')}

For mathematical content, use plain text format: write equations as "Z = c1*x1 + c2*x2 + ... + cn*xn" instead of LaTeX."""

    user_prompt = f"""Write comprehensive academic content for:

**Topic:** {section_info['section_number']} {section_info['section_title']}
**Course:** {course_context['course_title']}
**Description:** {section_info['description']}

Requirements:
- Include clear definitions with **bold key terms**
- Provide detailed explanations with examples
- Add case studies relevant to {course_context['target_audience']}
- Include practical applications
- Use *italics* for emphasis
- Write equations in plain text format (no LaTeX)
- Suggest image placements with descriptions

Make it engaging, informative, and academically rigorous."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    return make_api_call(messages, max_tokens=2500)

def generate_outline_with_ai():
    """
    Generate course outline using AI based on course configuration
    
    Returns:
        list: Outline with units and sections, or None if failed
    """
    num_units = st.session_state.get('num_units', 4)
    sections_per_unit = st.session_state.get('sections_per_unit', 8)
    
    st.write(f"🎯 Generating {num_units} units with {sections_per_unit} sections each...")
    
    # Build outcomes context
    outcomes_context = ""
    if st.session_state.program_objectives:
        outcomes_context += f"\n**Program Educational Objectives (PEO):** {st.session_state.program_objectives}"
    if st.session_state.program_outcomes:
        outcomes_context += f"\n**Program Outcomes (PO):** {st.session_state.program_outcomes}"
    if st.session_state.course_outcomes:
        outcomes_context += f"\n**Course Learning Outcomes (CO):** {st.session_state.course_outcomes}"
    if st.session_state.specialized_outcomes:
        outcomes_context += f"\n**Program Specific Outcomes (PSO):** {st.session_state.specialized_outcomes}"
    
    system_prompt = """You are an expert curriculum designer with deep knowledge of educational standards and course design.

Create a comprehensive, well-structured course outline that follows academic best practices.

CRITICAL: Output ONLY valid JSON in this EXACT format:
[
  {
    "unit_number": 1,
    "unit_title": "Descriptive, Engaging Unit Title",
    "sections": [
      {
        "section_number": "1.1",
        "section_title": "Clear Section Title",
        "description": "Detailed 2-3 sentence description of what this section covers, including key concepts and learning goals"
      }
    ]
  }
]

REQUIREMENTS:
- Create academically rigorous, engaging titles
- Make descriptions detailed and specific (not generic)
- Ensure logical progression and flow between topics
- Cover the subject comprehensively
- Use proper academic terminology
- Build complexity progressively from foundational to advanced
- Ensure each description is at least 2-3 sentences explaining WHAT will be covered"""

    user_prompt = f"""Create a comprehensive course outline for:

**Course Title:** {st.session_state.course_title}
**Course Code:** {st.session_state.course_code}
**Target Audience:** {st.session_state.target_audience}
**Credits:** {st.session_state.credits}

{outcomes_context}

**STRUCTURE REQUIREMENTS:**
- Generate EXACTLY {num_units} units
- Each unit must have EXACTLY {sections_per_unit} sections
- Section numbering: 1.1 to 1.{sections_per_unit}, 2.1 to 2.{sections_per_unit}, etc.
- Total sections must be: {num_units * sections_per_unit}

**CONTENT REQUIREMENTS:**
- Make unit titles specific to {st.session_state.course_title}
- Create logical flow from basics to advanced concepts
- Ensure comprehensive coverage of the subject
- Make descriptions specific and detailed (NOT generic like "First topic", "Overview")
- Each description should explain WHAT concepts/theories/applications will be covered

Return ONLY the JSON array. No additional text, no markdown code blocks, just pure JSON."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    with st.expander("🔍 AI Outline Generation Details", expanded=True):
        outline_str = make_api_call(messages, max_tokens=3000, retries=3)
    
    if outline_str:
        try:
            # Clean response
            outline_str = outline_str.strip()
            
            # Remove markdown code blocks if present
            json_match = re.search(r'\`\`\`(?:json)?\s*\n(.*?)\n\`\`\`', outline_str, re.DOTALL)
            if json_match:
                outline_str = json_match.group(1)
            
            outline_str = outline_str.strip()
            
            # Parse JSON
            parsed_outline = json.loads(outline_str)
            
            if isinstance(parsed_outline, list) and len(parsed_outline) > 0:
                actual_units = len(parsed_outline)
                actual_sections = sum(len(u.get('sections', [])) for u in parsed_outline)
                
                st.success(f"✅ AI generated {actual_units} units with {actual_sections} sections!")
                
                # Verify structure
                if actual_units != num_units:
                    st.warning(f"⚠️ Expected {num_units} units, got {actual_units}")
                
                if actual_sections < (num_units * sections_per_unit * 0.8):
                    st.warning(f"⚠️ Expected ~{num_units * sections_per_unit} sections, got {actual_sections}")
                
                # Show preview
                with st.expander("📋 Preview Generated Outline", expanded=False):
                    for unit in parsed_outline[:2]:
                        st.write(f"**Unit {unit['unit_number']}: {unit['unit_title']}**")
                        for section in unit.get('sections', [])[:3]:
                            st.write(f"  - {section['section_number']} {section['section_title']}")
                
                return parsed_outline
            else:
                st.error("❌ Invalid outline format - not a valid list")
                return None
                
        except json.JSONDecodeError as e:
            st.error(f"❌ JSON parsing failed: {str(e)}")
            st.write("**Raw AI Response:**")
            st.code(outline_str[:1000], language="json")
            st.error("💡 Try regenerating or check API response format")
            return None
    
    st.error("❌ AI did not return any content")
    return None

def generate_image_prompt_for_section(section_info, course_context):
    """
    Generate a detailed image generation prompt for a section
    
    Args:
        section_info: Dict with section details
        course_context: Dict with course context
        
    Returns:
        str: Image generation prompt
    """
    system_prompt = """You are an expert at creating detailed image generation prompts for educational content.

Create a detailed, specific prompt that can be used with AI image generators like DALL-E, Midjourney, or Stable Diffusion.

The prompt should:
- Be specific and descriptive
- Include style (e.g., "professional diagram", "infographic", "illustration")
- Mention colors if relevant
- Specify composition and layout
- Be suitable for educational materials
- Be 2-3 sentences maximum"""

    user_prompt = f"""Create an image generation prompt for this educational section:

**Section:** {section_info['section_number']} {section_info['section_title']}
**Course:** {course_context['course_title']}
**Description:** {section_info['description']}
**Audience:** {course_context['target_audience']}

Generate a detailed prompt for ONE relevant educational image."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    prompt = make_api_call(messages, max_tokens=200)
    if not prompt or not prompt.strip():
        prompt = f"Educational diagram about {section_info['section_title']}"
    return prompt

# ============================================================================
# PDF COMPILATION WITH IMAGES
# ============================================================================

def add_header_footer(canvas, doc, course_info, logo=None):
    """
    Add header and footer to PDF pages
    
    Args:
        canvas: ReportLab canvas
        doc: Document template
        course_info: Course information dict
        logo: Logo image file or None
    """
    canvas.saveState()
    
    # Header - Document heading
    heading = st.session_state.get('document_heading', course_info.get('course_title', ''))
    if heading:
        canvas.setFont("Helvetica-Bold", 11)
        canvas.drawString(72, doc.height + 50, heading[:80])
    
    # Logo in header (top right)
    if logo:
        try:
            logo.seek(0)
            logo_bytes = logo.read()
            logo_buffer = BytesIO(logo_bytes)
            img = PilImage.open(logo_buffer)
            
            # Save as temporary file for ReportLab
            temp_logo = BytesIO()
            img.save(temp_logo, format='PNG')
            temp_logo.seek(0)
            
            canvas.drawImage(
                temp_logo,
                doc.width - 1.5*inch + 72,
                doc.height + 30,
                width=1.2*inch,
                height=0.6*inch,
                preserveAspectRatio=True,
                mask='auto'
            )
        except Exception as e:
            pass  # Ignore logo errors
    
    # Footer - Page number and date
    canvas.setFont("Helvetica", 9)
    page_text = f"Page {doc.page}"
    canvas.drawString(72, 40, page_text)
    
    date_text = f"Generated: {datetime.now().strftime('%Y-%m-%d')}"
    canvas.drawString(doc.width - 72 - canvas.stringWidth(date_text, "Helvetica", 9), 40, date_text)
    
    canvas.restoreState()

def create_decorative_line():
    """Create decorative line table for PDF"""
    if not REPORTLAB_AVAILABLE:
        return None
        
    line_table = Table([['']], colWidths=[6.5*inch])
    line_table.setStyle(TableStyle([
        ('LINEABOVE', (0, 0), (-1, 0), 2, colors.HexColor('#1f77b4'))
    ]))
    return line_table

def compile_unit_pdf(unit_data, course_info, content_dict):
    """
    Compile a single unit into PDF with images and formatting
    
    Args:
        unit_data: Unit information dict
        course_info: Course information dict
        content_dict: Dictionary of generated content
        
    Returns:
        BytesIO: PDF buffer or None if failed
    """
    if not REPORTLAB_AVAILABLE:
        st.error("ReportLab not available")
        return None
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=1.2*inch,
        bottomMargin=1*inch
    )
    
    # Add header/footer callback
    def add_page_elements(canvas, doc):
        add_header_footer(canvas, doc, course_info, st.session_state.get('logo'))
    
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f77b4'),
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        spaceAfter=30,
        spaceBefore=20
    )
    
    section_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#2c3e50'),
        fontName='Helvetica-Bold',
        spaceAfter=12,
        spaceBefore=12
    )
    
    subsection_style = ParagraphStyle(
        'Subsection',
        parent=styles['Heading3'],
        fontSize=13,
        textColor=colors.HexColor('#34495e'),
        fontName='Helvetica-Bold',
        spaceAfter=10,
        spaceBefore=10
    )
    
    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['BodyText'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
        leading=14,
        fontName='Helvetica'
    )
    
    # Cover page
    story.append(Spacer(1, 1*inch))
    story.append(Paragraph(f"UNIT {unit_data['unit_number']}", styles['Heading2']))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(unit_data['unit_title'].upper(), title_style))
    story.append(Spacer(1, 0.5*inch))
    
    # Decorative line
    if create_decorative_line():
        story.append(create_decorative_line())
        story.append(Spacer(1, 0.3*inch))
    
    story.append(Paragraph(course_info.get('course_title', ''), styles['Heading3']))
    story.append(Paragraph(f"Course Code: {course_info.get('course_code', 'N/A')}", styles['Normal']))
    story.append(Paragraph(f"Credits: {course_info.get('credits', 3)}", styles['Normal']))
    story.append(PageBreak())
    
    # Content sections
    for section in unit_data.get('sections', []):
        sec_key = f"{section['section_number']} {section['section_title']}"
        
        # Section heading
        story.append(Paragraph(f"<b>{sec_key}</b>", section_style))
        if create_decorative_line():
            story.append(create_decorative_line())
        story.append(Spacer(1, 0.2*inch))
        
        # Add image if available
        image_data = st.session_state.images.get(sec_key)
        if image_data:
            try:
                image_data.seek(0)
                img_bytes = image_data.read()
                img_buffer = BytesIO(img_bytes)
                
                # Add image to PDF
                # Attempt to get image dimensions for better scaling
                try:
                    img_pil = PilImage.open(img_buffer)
                    img_width, img_height = img_pil.size
                    aspect_ratio = img_height / img_width
                    max_width = 4 * inch
                    max_height = 2.5 * inch
                    
                    if img_width > max_width:
                        draw_width = max_width
                        draw_height = max_width * aspect_ratio
                        if draw_height > max_height:
                            draw_height = max_height
                            draw_width = max_height / aspect_ratio

                    else:
                        draw_width = img_width
                        draw_height = img_height
                        if draw_height > max_height:
                            draw_height = max_height
                            draw_width = max_height / aspect_ratio

                except Exception as pil_err:
                    st.warning(f"PIL error for {sec_key}: {pil_err}. Using default image size.")
                    draw_width = 4 * inch
                    draw_height = 2.5 * inch
                
                img_buffer.seek(0) # Reset buffer after reading with PIL
                story.append(Image(img_buffer, width=draw_width, height=draw_height))
                story.append(Spacer(1, 0.2*inch))
            except Exception as e:
                st.warning(f"⚠️ Could not add image for {sec_key}: {str(e)}")
        
        # Content
        content = content_dict.get(sec_key, "[Content not generated]")
        content_lines = content.split('\n')
        
        for line in content_lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            # Handle different content types
            if 'CHECK YOUR PROGRESS' in line.upper():
                story.append(Paragraph("<b>CHECK YOUR PROGRESS</b>", subsection_style))
                story.append(Spacer(1, 0.1*inch))
            elif line.startswith('## '):
                clean_line = clean_text_for_pdf(line[3:])
                story.append(Paragraph(clean_line, subsection_style))
            elif line.startswith(('*', '-', '•', '1.', '2.', '3.')):
                clean_line = clean_text_for_pdf(re.sub(r'^[\*\-•\d\.]\s*', '', line))
                # Use bullet style for lists
                style = ParagraphStyle(
                    'ListItem',
                    parent=body_style,
                    bulletIndent=18,
                    leftIndent=36,
                    spaceBefore=6,
                    spaceAfter=6,
                    bulletFontName='Symbol', # Use a standard symbol font for bullets
                    bulletText='•'
                )
                story.append(Paragraph(clean_line, style))
            else:
                clean_line = clean_text_for_pdf(line)
                if len(clean_line) > 3:
                    try:
                        story.append(Paragraph(clean_line, body_style))
                    except Exception as e:
                        # Fallback for problematic content
                        st.warning(f"⚠️ Problematic paragraph formatting for '{line[:50]}...': {e}. Using raw text.")
                        story.append(Paragraph(line[:500], body_style)) # Limit raw text to avoid issues
        
        # Add a page break after each section for better readability, but only if not the last section of the unit
        if section != unit_data.get('sections', [])[-1]:
            story.append(PageBreak())
        else: # Add space after the last section of the unit
            story.append(Spacer(1, 0.5*inch))
    
    try:
        doc.build(story, onFirstPage=add_page_elements, onLaterPages=add_page_elements)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ PDF compilation error: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None

def compile_complete_pdf(outline, course_info, content_dict):
    """
    Compile complete course PDF with all units
    
    Args:
        outline: Complete course outline
        course_info: Course information
        content_dict: All generated content
        
    Returns:
        BytesIO: Complete PDF buffer or None
    """
    if not REPORTLAB_AVAILABLE:
        st.error("ReportLab not available")
        return None
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=1.2*inch,
        bottomMargin=1*inch
    )
    
    def add_page_elements(canvas, doc):
        add_header_footer(canvas, doc, course_info, st.session_state.get('logo'))
    
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        fontSize=28,
        alignment=TA_CENTER,
        spaceAfter=40,
        fontName='Helvetica-Bold'
    )
    
    # Title page
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph(course_info['course_title'], title_style))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(
        f"Course Code: {course_info['course_code']} | Credits: {course_info['credits']}",
        styles['Normal']
    ))
    story.append(PageBreak())
    
    # Table of Contents Placeholder (Optional - can be complex to generate dynamically)
    # story.append(Paragraph("Table of Contents", styles['Heading1']))
    # story.append(PageBreak())
    
    # All units content
    for unit in outline:
        # Unit Title Page
        unit_title_style = ParagraphStyle(
            'UnitTitle',
            parent=styles['Heading1'],
            fontSize=22,
            alignment=TA_CENTER,
            spaceAfter=30,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#1f77b4')
        )
        story.append(Paragraph(f"UNIT {unit['unit_number']}", styles['Heading2']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(unit['unit_title'], unit_title_style))
        story.append(PageBreak())
        
        for section in unit.get('sections', []):
            sec_key = f"{section['section_number']} {section['section_title']}"
            story.append(Paragraph(sec_key, styles['Heading2']))
            
            # Add image if available
            image_data = st.session_state.images.get(sec_key)
            if image_data:
                try:
                    image_data.seek(0)
                    img_bytes = image_data.read()
                    img_buffer = BytesIO(img_bytes)
                    
                    try:
                        img_pil = PilImage.open(img_buffer)
                        img_width, img_height = img_pil.size
                        aspect_ratio = img_height / img_width
                        max_width = 4 * inch
                        max_height = 2.5 * inch
                        
                        if img_width > max_width:
                            draw_width = max_width
                            draw_height = max_width * aspect_ratio
                            if draw_height > max_height:
                                draw_height = max_height
                                draw_width = max_height / aspect_ratio
                        else:
                            draw_width = img_width
                            draw_height = img_height
                            if draw_height > max_height:
                                draw_height = max_height
                                draw_width = max_height / aspect_ratio

                    except Exception as pil_err:
                        st.warning(f"PIL error for {sec_key}: {pil_err}. Using default image size.")
                        draw_width = 4 * inch
                        draw_height = 2.5 * inch
                    
                    img_buffer.seek(0)
                    story.append(Image(img_buffer, width=draw_width, height=draw_height))
                    story.append(Spacer(1, 0.2*inch))
                except:
                    st.warning(f"Could not add image to complete PDF for {sec_key}")
            
            content = content_dict.get(sec_key, "[Not generated]")
            clean_content = clean_text_for_pdf(content)
            
            # Add content paragraphs
            current_paragraph_text = ""
            for line in clean_content.split('\n'):
                line = line.strip()
                if not line:
                    if current_paragraph_text:
                        try:
                            story.append(Paragraph(current_paragraph_text, styles['Normal']))
                        except Exception as e:
                            st.warning(f"Error formatting paragraph for '{current_paragraph_text[:50]}...': {e}")
                            story.append(Paragraph(current_paragraph_text[:500], styles['Normal']))
                        current_paragraph_text = ""
                    story.append(Spacer(1, 6))
                    continue

                # Basic markdown list handling for complete PDF (less detailed than unit)
                if line.startswith(('*', '-', '•')) and current_paragraph_text:
                    # If we have pending paragraph text, finish it first
                    try:
                        story.append(Paragraph(current_paragraph_text, styles['Normal']))
                    except Exception as e:
                        st.warning(f"Error formatting paragraph for '{current_paragraph_text[:50]}...': {e}")
                        story.append(Paragraph(current_paragraph_text[:500], styles['Normal']))
                    current_paragraph_text = ""
                    
                    # Start new list item paragraph
                    clean_line = clean_text_for_pdf(re.sub(r'^[\*\-•\d\.]\s*', '', line))
                    style = ParagraphStyle(
                        'ListItemComplete',
                        parent=styles['Normal'],
                        bulletIndent=18,
                        leftIndent=36,
                        spaceBefore=6,
                        spaceAfter=6,
                        bulletFontName='Symbol',
                        bulletText='•'
                    )
                    try:
                        story.append(Paragraph(clean_line, style))
                    except Exception as e:
                        st.warning(f"Error formatting list item for '{clean_line[:50]}...': {e}")
                        story.append(Paragraph(clean_line[:500], styles['Normal']))
                
                else:
                    # Append to current paragraph or start new one
                    if current_paragraph_text:
                        current_paragraph_text += " " + line # Append with space
                    else:
                        current_paragraph_text = line
            
            # Add any remaining paragraph text
            if current_paragraph_text:
                try:
                    story.append(Paragraph(current_paragraph_text, styles['Normal']))
                except Exception as e:
                    st.warning(f"Error formatting final paragraph for '{current_paragraph_text[:50]}...': {e}")
                    story.append(Paragraph(current_paragraph_text[:500], styles['Normal']))
            
            story.append(PageBreak())
    
    try:
        doc.build(story, onFirstPage=add_page_elements, onLaterPages=add_page_elements)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Complete PDF error: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None

# ============================================================================
# DOCX COMPILATION WITH IMAGES
# ============================================================================

def compile_unit_docx(unit_data, course_info, content_dict):
    """
    Compile unit as editable DOCX with images
    
    Args:
        unit_data: Unit information
        course_info: Course information
        content_dict: Generated content
        
    Returns:
        BytesIO: DOCX buffer or None
    """
    if not DOCX_AVAILABLE:
        st.error("python-docx not available")
        return None
    
    buffer = BytesIO()
    doc = Document()
    
    # Custom styles (optional, but good practice)
    try:
        styles = doc.styles
        if 'Heading 1' not in styles:
             styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        if 'Heading 2' not in styles:
             styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        
        h1 = styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(16)
        h1.font.bold = True
        
        h2 = styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(13)
        h2.font.bold = True
        
        normal = styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(11)

    except Exception as e:
        st.warning(f"Could not apply custom styles to DOCX: {e}")

    # Title
    title = doc.add_heading(f"UNIT {unit_data['unit_number']}: {unit_data['unit_title']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Course info
    p_course_info = doc.add_paragraph()
    p_course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_course_info.add_run(f"Course: {course_info.get('course_title', 'N/A')} | ").font.size = Pt(11)
    p_course_info.add_run(f"Code: {course_info.get('course_code', 'N/A')} | ").font.size = Pt(11)
    p_course_info.add_run(f"Credits: {course_info.get('credits', 3)}").font.size = Pt(11)
    
    doc.add_page_break()
    
    # Content sections
    for section in unit_data.get('sections', []):
        sec_key = f"{section['section_number']} {section['section_title']}"
        
        # Section heading
        doc.add_heading(sec_key, level=1)
        
        # Add image if available
        image_data = st.session_state.images.get(sec_key)
        if image_data:
            try:
                image_data.seek(0)
                # Attempt to get image dimensions for better scaling
                try:
                    img_pil = PilImage.open(BytesIO(image_data.read()))
                    img_width, img_height = img_pil.size
                    aspect_ratio = img_height / img_width
                    max_width = 4.5 # Inches
                    
                    if img_width > max_width * 72: # DPI
                        draw_width = Inches(max_width)
                        draw_height = draw_width * aspect_ratio
                    else:
                        draw_width = Inches(img_width / 72) # Convert pixels to inches
                        draw_height = Inches(img_height / 72)
                        if draw_height > Inches(3.0): # Limit height
                            draw_height = Inches(3.0)
                            draw_width = draw_height / aspect_ratio

                except Exception as pil_err:
                    st.warning(f"PIL error for DOCX image {sec_key}: {pil_err}. Using default image size.")
                    draw_width = Inches(4.5)
                    draw_height = Inches(3.0)
                
                image_data.seek(0) # Reset buffer
                doc.add_picture(image_data, width=draw_width)
                doc.add_paragraph()  # Space after image
            except Exception as e:
                st.warning(f"⚠️ Could not add image to DOCX: {str(e)}")
        
        # Add content
        content = content_dict.get(sec_key, "[Not generated]")
        
        # Process content line by line to handle basic markdown
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph() # Add empty paragraph for spacing
                continue
            
            p = doc.add_paragraph()
            
            # Basic handling for bold (**) and italics (*)
            parts = re.split(r'(\*\*.+?\*\*|\*.+?\*)', line)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
        
        doc.add_page_break()
    
    try:
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ DOCX error: {str(e)}")
        return None

def compile_complete_docx(outline, course_info, content_dict):
    """
    Compile complete course as DOCX
    
    Args:
        outline: Course outline
        course_info: Course information
        content_dict: All content
        
    Returns:
        BytesIO: DOCX buffer or None
    """
    if not DOCX_AVAILABLE:
        return None
    
    buffer = BytesIO()
    doc = Document()
    
    # Custom styles
    try:
        styles = doc.styles
        if 'Heading 1' not in styles: styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        if 'Heading 2' not in styles: styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        
        h1 = styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(16)
        h1.font.bold = True
        
        h2 = styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(13)
        h2.font.bold = True
        
        normal = styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(11)

    except Exception as e:
        st.warning(f"Could not apply custom styles to DOCX: {e}")

    # Title page
    title = doc.add_heading(course_info['course_title'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_course_info = doc.add_paragraph()
    p_course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_course_info.add_run(f"Code: {course_info.get('course_code', 'N/A')} | ").font.size = Pt(11)
    p_course_info.add_run(f"Credits: {course_info.get('credits', 3)}").font.size = Pt(11)
    
    doc.add_page_break()
    
    # All units
    for unit in outline:
        doc.add_heading(f"UNIT {unit['unit_number']}: {unit['unit_title']}", level=1)
        
        for section in unit.get('sections', []):
            sec_key = f"{section['section_number']} {section['section_title']}"
            doc.add_heading(sec_key, level=2)
            
            # Add image
            image_data = st.session_state.images.get(sec_key)
            if image_data:
                try:
                    image_data.seek(0)
                    # Attempt to get image dimensions for better scaling
                    try:
                        img_pil = PilImage.open(BytesIO(image_data.read()))
                        img_width, img_height = img_pil.size
                        aspect_ratio = img_height / img_width
                        max_width = 4.5 # Inches
                        
                        if img_width > max_width * 72: # DPI
                            draw_width = Inches(max_width)
                            draw_height = draw_width * aspect_ratio
                        else:
                            draw_width = Inches(img_width / 72)
                            draw_height = Inches(img_height / 72)
                            if draw_height > Inches(3.0):
                                draw_height = Inches(3.0)
                                draw_width = draw_height / aspect_ratio

                    except Exception as pil_err:
                        st.warning(f"PIL error for DOCX image {sec_key}: {pil_err}. Using default image size.")
                        draw_width = Inches(4.5)
                        draw_height = Inches(3.0)
                    
                    image_data.seek(0)
                    doc.add_picture(image_data, width=draw_width)
                except Exception as e:
                    st.warning(f"Could not add image to complete DOCX: {str(e)}")
            
            # Add content
            content = content_dict.get(sec_key, "[Not generated]")
            
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    # Simple markdown handling for bold and italic in complete DOCX
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.+?\*\*|\*.+?\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        elif part.startswith('*') and part.endswith('*'):
                            run = p.add_run(part[1:-1])
                            run.italic = True
                        else:
                            p.add_run(part)
            doc.add_page_break()
    
    try:
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Complete DOCX error: {str(e)}")
        return None

# ============================================================================
# Phase 3 Complete
# ============================================================================
print("Phase 3: Content Generation and Compilation loaded successfully")
"""
"""

# ============================================================================
# NAVIGATION
# ============================================================================

def show_navigation():
    """Display step navigation bar"""
    steps = {
        'syllabus_upload': '1️⃣ Syllabus',
        'configuration': '2️⃣ Config',
        'outline_generation': '3️⃣ Outline',
        'content_generation': '4️⃣ Content',
        'compilation': '5️⃣ Output'
    }
    
    current = st.session_state.step
    
    cols = st.columns(len(steps))
    for i, (key, name) in enumerate(steps.items()):
        with cols[i]:
            if key == current:
                st.markdown(f"**🔵 {name}**")
            else:
                st.markdown(f"⚪ {name}")
    
    st.divider()

# ============================================================================
# PAGE 1: SYLLABUS UPLOAD
# ============================================================================

def show_syllabus_upload_page():
    """Syllabus upload page - optional PDF upload"""
    st.header("📄 Step 1: Syllabus Upload (Optional)")
    
    st.info("💡 You can upload a syllabus PDF to auto-extract structure, or skip this step to create a custom outline with AI.")
    
    choice = st.radio(
        "Choose an option:",
        ["Upload Syllabus PDF", "Skip and Create Custom Outline"],
        key="upload_choice"
    )
    
    if choice == "Upload Syllabus PDF":
        if PYPDF2_AVAILABLE:
            uploaded = st.file_uploader(
                "Upload PDF Syllabus",
                type=['pdf'],
                key="syllabus_file",
                help="Upload your course syllabus PDF to automatically extract the structure"
            )
            
            if uploaded:
                with st.spinner("📖 Extracting text from PDF..."):
                    text = extract_pdf_text(uploaded)
                    
                    if text:
                        st.success("✅ Text extracted successfully!")
                        
                        with st.expander("📝 Preview Extracted Text", expanded=False):
                            st.text_area("Extracted Content", text[:1000] + "...", height=200)
                        
                        with st.spinner("🔍 Parsing syllabus structure..."):
                            structure = parse_syllabus_structure(text)
                            st.session_state.extracted_structure = structure
                            
                            if structure['course_info']:
                                st.success("✅ Course information extracted!")
                                with st.expander("📋 Course Information", expanded=True):
                                    for key, value in structure['course_info'].items():
                                        st.write(f"**{key.title()}:** {value}")
                            
                            if structure['units']:
                                st.success(f"✅ Found {len(structure['units'])} units")
                                
                                for unit in structure['units']:
                                    with st.expander(f"Unit {unit['unit_number']}: {unit['unit_title']}", expanded=False):
                                        st.write(f"**Topics ({len(unit['topics'])}):**")
                                        for i, topic in enumerate(unit['topics'], 1):
                                            st.write(f"{i}. {topic}")
                                
                                st.divider()
                                
                                if st.checkbox("✅ Structure looks good, use this", key="confirm_structure"):
                                    if st.button("Continue to Configuration →", type="primary", key="continue_with_structure"):
                                        st.session_state.step = 'configuration'
                                        st.rerun()
                            else:
                                st.warning("⚠️ No units found in syllabus. You can still continue and create a custom outline.")
                                if st.button("Continue Anyway →", type="primary"):
                                    st.session_state.step = 'configuration'
                                    st.rerun()
        else:
            st.error("❌ PyPDF2 not installed. Cannot extract PDF content.")
            st.info("Install with: pip install PyPDF2")
            if st.button("Continue Without Upload →", type="primary"):
                st.session_state.step = 'configuration'
                st.rerun()
    else:
        st.info("✅ You'll create a custom outline in the next steps using AI")
        if st.button("Continue to Configuration →", type="primary", key="skip_upload"):
            st.session_state.step = 'configuration'
            st.rerun()

# ============================================================================
# PAGE 2: CONFIGURATION
# ============================================================================

def show_configuration_page():
    """Configuration page for course details and PDF upload"""
    st.header("⚙️ Step 2: Configuration")
    
    # ========== API Configuration ==========
    st.subheader("🔑 API Configuration")
    
    col1, col2 = st.columns([3, 1])
    with col1:
        api_key = st.text_input(
            "Grok API Key",
            value=st.session_state.api_key,
            type="password",
            key="api_key_input",
            help="Your Grok API key from x.ai - starts with 'xai-'"
        )
        st.session_state.api_key = api_key
        
        if api_key and api_key.startswith('xai-'):
            st.success("✅ Valid API key format")
        else:
            st.warning("⚠️ API key should start with 'xai-'")
    
    with col2:
        st.write("")
        st.write("")
        if st.button("🧪 Test API", use_container_width=True, key="test_api_btn"):
            with st.expander("🔍 API Test Results", expanded=True):
                st.info("Testing API connection and response quality...")
                
                test_messages = [
                    {"role": "system", "content": "You are a helpful academic content generator."},
                    {"role": "user", "content": "Write a 200-word introduction about Organizational Behaviour. Include: definition, importance, and key concepts. Use academic tone."}
                ]
                
                test_response = make_api_call(test_messages, max_tokens=500)
                
                if test_response:
                    st.success("✅ API is working!")
                    st.write("**Response Preview:**")
                    st.write(test_response[:300] + "...")
                    
                    word_count = len(test_response.split())
                    if word_count >= 150:
                        st.success(f"✅ Good response length: {word_count} words")
                    else:
                        st.warning(f"⚠️ Response seems short: {word_count} words")
                    
                    if len(test_response) > 100:
                        st.success("✅ API returning substantial content")
                        st.info("💡 Your API is ready for curriculum generation!")
                    else:
                        st.error("❌ API response too short - check configuration")
                else:
                    st.error("❌ API test failed - check logs above")
    
    st.divider()
    
    # ========== Course Details ==========
    st.subheader("📚 Course Details")
    
    extracted = st.session_state.get('extracted_structure') or {}
    course_info_from_syllabus = extracted.get('course_info', {})
    
    col1, col2 = st.columns(2)
    with col1:
        title = st.text_input(
            "Course Title",
            value=course_info_from_syllabus.get('title', st.session_state.course_title),
            key="title_input",
            help="Full name of the course"
        )
        st.session_state.course_title = title
        
        code = st.text_input(
            "Course Code",
            value=course_info_from_syllabus.get('code', st.session_state.course_code),
            key="code_input",
            help="Course code (e.g., MBA101, CS202)"
        )
        st.session_state.course_code = code
    
    with col2:
        credits = st.number_input(
            "Credits",
            min_value=1,
            max_value=10,
            value=int(course_info_from_syllabus.get('credits', st.session_state.credits)) if course_info_from_syllabus.get('credits') else st.session_state.credits,
            key="credits_input",
            help="Number of credit hours"
        )
        st.session_state.credits = credits
        
        audience = st.selectbox(
            "Target Audience",
            ["Postgraduate (MBA)", "Undergraduate", "Diploma", "Certificate"],
            index=["Postgraduate (MBA)", "Undergraduate", "Diploma", "Certificate"].index(st.session_state.target_audience) if st.session_state.target_audience in ["Postgraduate (MBA)", "Undergraduate", "Diploma", "Certificate"] else 0,
            key="audience_select",
            help="Target student level"
        )
        st.session_state.target_audience = audience
    
    # Add new fields to Course Details
    st.info("💡 Provide additional course details (optional, but recommended for better content).")
    col1, col2 = st.columns(2)
    with col1:
        learning_objectives = st.text_area(
            "Learning Objectives (Course)",
            value=st.session_state.learning_objectives,
            placeholder="e.g., Understand fundamental concepts of X, Analyze Y using Z models, Apply theory A to case study B",
            height=120,
            key="learning_objectives_input"
        )
        st.session_state.learning_objectives = learning_objectives
        
        prerequisites = st.text_area(
            "Prerequisites",
            value=st.session_state.prerequisites,
            placeholder="e.g., Basic understanding of statistics, Completion of MATH101",
            height=120,
            key="prerequisites_input"
        )
        st.session_state.prerequisites = prerequisites
    
    with col2:
        assessment_methods = st.text_area(
            "Assessment Methods",
            value=st.session_state.assessment_methods,
            placeholder="e.g., Mid-term exam (30%), Final exam (40%), Assignments (20%), Project (10%)",
            height=120,
            key="assessment_methods_input"
        )
        st.session_state.assessment_methods = assessment_methods
        
        # Placeholder for future features like Grading Policy
        st.text_area(
            "Grading Policy (Future Feature)",
            disabled=True,
            height=120,
            help="Details about grading scales and policies will be added here."
        )
    
    st.divider()
    
    # ========== Document Customization ==========
    st.subheader("📄 Document Customization")
    
    col1, col2 = st.columns(2)
    with col1:
        document_heading = st.text_input(
            "Document Header Text (Optional)",
            value=st.session_state.document_heading,
            key="doc_heading",
            help="Text to appear in the header of each page",
            placeholder="e.g., XYZ University - MBA Program"
        )
        st.session_state.document_heading = document_heading
    
    with col2:
        logo = st.file_uploader(
            "Upload Logo for Header (Optional)",
            type=['png', 'jpg', 'jpeg'],
            key="logo_uploader",
            help="Logo will appear in the top-right corner of each page"
        )
        if logo:
            st.session_state.logo = logo
            st.success("✅ Logo uploaded")
            # Show preview
            st.image(logo, width=150, caption="Logo Preview")
        elif st.session_state.logo:
            st.info("✅ Logo from previous session")
    
    st.divider()
    
    # ========== Content Structure ==========
    if not st.session_state.get('extracted_structure'):
        st.subheader("📚 Content Structure")
        st.info("💡 Define how many units and sections you want. AI will generate relevant content for each.")
        
        col1, col2 = st.columns(2)
        with col1:
            num_units = st.number_input(
                "Number of Units",
                min_value=1,
                max_value=10,
                value=st.session_state.num_units,
                help="How many major units/modules for this course",
                key="num_units_config"
            )
            st.session_state.num_units = num_units
        
        with col2:
            sections_per_unit = st.number_input(
                "Sections per Unit",
                min_value=3,
                max_value=15,
                value=st.session_state.sections_per_unit,
                help="How many topics/sections in each unit",
                key="sections_per_unit_config"
            )
            st.session_state.sections_per_unit = sections_per_unit
        
        total_sections = num_units * sections_per_unit
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Units", num_units)
        with col2:
            st.metric("Total Sections", total_sections)
        with col3:
            st.metric("Est. Pages", f"~{total_sections * 5}")
        
        st.caption(f"💡 AI will generate {num_units} units with {sections_per_unit} topics each = {total_sections} total sections")
    else:
        st.info("✅ Using structure from uploaded syllabus")
    
    st.divider()
    
    # ========== Academic Mappings ==========
    st.subheader("🎯 Academic Mappings (Optional but Recommended)")
    
    with st.expander("ℹ️ What are PEO, PO, CO, PSO?", expanded=False):
        st.markdown("""
        **Program Educational Objectives (PEO):**
        - What students can accomplish after completing the entire program
        - Long-term career and professional achievements
        
        **Program Outcomes (PO):**
        - Skills and knowledge students gain from the program
        - Example: PO1: Critical thinking, PO2: Communication skills, PO3: Ethical awareness
        
        **Course Learning Outcomes (CO):**
        - What students learn in THIS specific course
        - Example: CO1: Understand OB concepts [Bloom: Understand], CO2: Apply leadership theories [Bloom: Apply]
        
        **Program Specific Outcomes (PSO):**
        - Specialized skills specific to the program (e.g., MBA-specific)
        - Example: PSO1: Strategic management, PSO2: Business analytics
        
        These mappings help create better-aligned content and demonstrate learning outcomes.
        """)
    
    st.info("💡 These mappings enhance content quality and alignment. Leave blank if not needed.")
    
    peo = st.text_area(
        "Program Educational Objectives (PEO)",
        value=st.session_state.program_objectives,
        placeholder="Example:\n- Develop strategic leadership capabilities\n- Foster analytical decision-making skills\n- Build effective communication abilities",
        help="What students should achieve after completing the program",
        key="peo_input",
        height=120
    )
    st.session_state.program_objectives = peo
    
    col1, col2 = st.columns(2)
    with col1:
        po = st.text_area(
            "Program Outcomes (PO)",
            value=st.session_state.program_outcomes,
            placeholder="Example:\nPO1: Critical thinking and problem-solving\nPO2: Effective communication\nPO3: Ethical decision-making\nPO4: Teamwork and collaboration",
            help="Skills and knowledge from the program",
            key="po_input",
            height=150
        )
        st.session_state.program_outcomes = po
    
    with col2:
        pso = st.text_area(
            "Program Specific Outcomes (PSO)",
            value=st.session_state.specialized_outcomes,
            placeholder="Example:\nPSO1: Advanced managerial skills\nPSO2: Strategic HR management\nPSO3: Organizational leadership\nPSO4: Change management expertise",
            help="Specialized skills for this specific program",
            key="pso_input",
            height=150
        )
        st.session_state.specialized_outcomes = pso
    
    co = st.text_area(
        "Course Learning Outcomes (CO)",
        value=st.session_state.course_outcomes,
        placeholder="Example:\nCO1: Understand key organizational behaviour concepts [Bloom: Understand]\nCO2: Apply OB theories to real-world scenarios [Bloom: Apply]\nCO3: Analyze organizational dynamics and culture [Bloom: Analyze]\nCO4: Evaluate organizational strategies [Bloom: Evaluate]",
        help="What students will learn in THIS specific course",
        key="co_input",
        height=150
    )
    st.session_state.course_outcomes = co
    
    if peo or po or co or pso:
        st.success("✅ Academic mappings will be integrated into content generation")
    else:
        st.info("ℹ️ Content will be generated with general academic outcomes")
    
    # Navigation
    st.divider()
    if st.button("Next: Generate Outline →", type="primary", use_container_width=True, key="config_next"):
        if not st.session_state.course_title:
            st.error("❌ Please enter a course title")
        else:
            st.session_state.step = 'outline_generation'
            st.rerun()

# ============================================================================
# PAGE 3: OUTLINE GENERATION
# ============================================================================

def show_outline_page():
    """Outline generation and editing page"""
    st.header("📋 Step 3: Course Outline")
    
    # Check if outline needs to be generated
    if not st.session_state.outline_generated or not st.session_state.approved_outline:
        
        # Check if extracted from syllabus
        extracted = st.session_state.get('extracted_structure')
        
        if extracted and extracted.get('units'):
            st.info("✅ Using syllabus structure from uploaded PDF")
            
            # Convert extracted structure to outline format
            outline = []
            for unit in extracted['units']:
                sections = []
                for i, topic in enumerate(unit['topics'], 1):
                    sections.append({
                        "section_number": f"{unit['unit_number']}.{i}",
                        "section_title": topic,
                        "description": topic
                    })
                
                outline.append({
                    "unit_number": unit['unit_number'],
                    "unit_title": unit['unit_title'],
                    "sections": sections
                })
            
            st.session_state.approved_outline = outline
            st.session_state.raw_outline = outline # Store as raw outline
            st.session_state.outline_generated = True
            st.success(f"✅ Created outline with {len(outline)} units")
            st.rerun()
            
        else:
            # MUST generate with AI - NO DEFAULTS
            st.warning("⚠️ No outline generated yet")
            st.info("💡 Click 'Generate with AI' to create a custom course outline based on your configuration")
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("🤖 Generate Outline with AI", type="primary", use_container_width=True, key="generate_ai_outline"):
                    with st.spinner("🤖 AI is creating your course outline... This may take 30-60 seconds"):
                        generated_outline = generate_outline_with_ai()
                        
                        if generated_outline:
                            st.session_state.raw_outline = generated_outline # Store as raw outline
                            st.session_state.approved_outline = generated_outline # Also set as approved initially
                            st.session_state.outline_generated = True
                            st.success("✅ Outline generated successfully!")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("❌ AI generation failed. Please try again or check your API key.")
                            return
            
            with col2:
                if st.button("← Back to Configuration", use_container_width=True, key="back_no_outline"):
                    st.session_state.step = 'configuration'
                    st.rerun()
            
            return
    
    # Display and edit outline
    outline = st.session_state.approved_outline
    total_sections = sum(len(u.get('sections', [])) for u in outline)
    
    # Summary metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📚 Units", len(outline))
    with col2:
        st.metric("📝 Sections", total_sections)
    with col3:
        st.metric("📄 Est. Pages", f"~{total_sections * 5}")
    
    st.divider()
    
    # Convert outline to editable table format
    st.subheader("✏️ Review and Edit Outline")
    st.caption("Click any cell to edit directly. You can modify titles and descriptions.")
    
    rows = []
    for unit in outline:
        for section in unit.get('sections', []):
            rows.append({
                'Unit': unit['unit_number'],
                'Unit Title': unit['unit_title'],
                'Section': section['section_number'],
                'Section Title': section['section_title'],
                'Description': section.get('description', '') # Use .get for safety
            })
    
    edited = st.data_editor(
        rows,
        num_rows="dynamic",
        use_container_width=True,
        height=400,
        key="outline_editor",
        column_config={
            "Unit": st.column_config.NumberColumn("Unit #", width="small"),
            "Unit Title": st.column_config.TextColumn("Unit Title", width="medium"),
            "Section": st.column_config.TextColumn("Section #", width="small"),
            "Section Title": st.column_config.TextColumn("Section Title", width="medium"),
            "Description": st.column_config.TextColumn("Description", width="large"),
        }
    )
    
    st.divider()
    
    # Navigation buttons
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("← Back", use_container_width=True, key="back_from_outline"):
            st.session_state.step = 'configuration'
            st.rerun()
    
    with col2:
        if st.button("🔄 Regenerate with AI", use_container_width=True, key="regen_outline_btn"):
            st.session_state.raw_outline = ''
            st.session_state.approved_outline = []
            st.session_state.outline_generated = False
            st.rerun()
    
    with col3:
        if st.button("✅ Approve & Generate Content →", type="primary", use_container_width=True, key="approve_outline_btn"):
            # Convert edited data back to outline format
            approved = []
            current_unit = None
            
            # Sort rows by Unit and then Section for proper reconstruction
            sorted_edited = sorted(edited, key=lambda x: (int(x['Unit']), x['Section']))

            for row in sorted_edited:
                unit_num = int(row['Unit'])
                unit_title = row['Unit Title']
                section_num = row['Section']
                section_title = row['Section Title']
                description = row.get('Description', '') # Use .get for safety
                
                if current_unit is None or current_unit['unit_number'] != unit_num:
                    if current_unit:
                        approved.append(current_unit)
                    current_unit = {
                        'unit_number': unit_num,
                        'unit_title': unit_title,
                        'sections': []
                    }
                
                current_unit['sections'].append({
                    'section_number': section_num,
                    'section_title': section_title,
                    'description': description
                })
            
            if current_unit: # Add the last processed unit
                approved.append(current_unit)
            
            st.session_state.approved_outline = approved
            st.session_state.images = {}  # Reset images
            st.session_state.image_prompts = {}  # Reset image prompts
            st.session_state.content = {} # Clear previous content if re-generating
            st.session_state.content_status = {} # Clear previous content status
            st.session_state.step = 'content_generation'
            st.success("✅ Outline approved! Moving to content generation...")
            time.sleep(1)
            st.rerun()

# ============================================================================
# Phase 4 Part 1 Complete
# ============================================================================
print("Phase 4 Part 1: UI Pages (Syllabus, Config, Outline) loaded successfully")
"""
PHASE 4: USER INTERFACE PAGES (PART 2)
========================================
- Content generation page with image uploads
- Compilation page with download options
- Sidebar status display
"""

# ============================================================================
# SIDEBAR STATUS
# ============================================================================

def show_sidebar_status():
    """Enhanced sidebar with navigation and status"""
    with st.sidebar:
        st.title("🎓 AI Curriculum Generator")
        
        # Status indicators
        st.caption(f"PDF: {'✅' if REPORTLAB_AVAILABLE else '❌'}")
        st.caption(f"DOCX: {'✅' if DOCX_AVAILABLE else '❌'}")
        
        st.divider()

        st.header("Current Project Status")
        
        # Display current step
        steps = {
            'syllabus_upload': '1️⃣ Syllabus',
            'configuration': '2️⃣ Config',
            'outline_generation': '3️⃣ Outline',
            'content_generation': '4️⃣ Content',
            'compilation': '5️⃣ Output'
        }
        current_step_name = steps.get(st.session_state.step, "Unknown")
        st.write(f"**Current Step:** {current_step_name}")
        
        if st.session_state.step == 'configuration':
            pass # No specific status to show here, rely on config page
        elif st.session_state.step == 'outline_generation':
            outline_status = "✅ Generated" if st.session_state.outline_generated else "❌ Not Generated"
            st.write(f"Outline: {outline_status}")
        elif st.session_state.step == 'content_generation':
            total_sections = len(st.session_state.sections_to_process) if st.session_state.sections_to_process else 0
            completed_sections = len(st.session_state.content)
            progress_pct = (completed_sections / total_sections * 100) if total_sections > 0 else 0
            st.progress(progress_pct / 100)
            st.write(f"Content: {completed_sections}/{total_sections} sections ({progress_pct:.0f}%)")
            if st.session_state.paused:
                st.warning("Generation is paused")
        elif st.session_state.step == 'compilation':
            if st.session_state.compiled_files:
                num_files = len(st.session_state.compiled_files)
                st.success(f"{num_files} file(s) ready for download.")
            else:
                st.info("Compilation not started yet.")

        st.divider()

        st.caption("© 2024 AI Curriculum Generator")


# ============================================================================
# PAGE 4: CONTENT GENERATION WITH IMAGE UPLOADS
# ============================================================================

def show_content_generation_page():
    """Content generation page with image upload and prompt generation"""
    st.header("✍️ Step 4: Content Generation")
    
    if 'approved_outline' not in st.session_state or not st.session_state.approved_outline:
        st.error("❌ No approved outline found")
        if st.button("← Back to Outline", key="back_no_outline_gen"):
            st.session_state.step = 'outline_generation'
            st.rerun()
        return
    
    # Initialize generation if not already started or if resuming
    if not st.session_state.content or not st.session_state.content_status:
        st.session_state.content = {}
        st.session_state.content_status = {} # Store status per section
        st.session_state.images = {}
        st.session_state.image_prompts = {}
        st.session_state.generation_start_time = time.time()
        st.session_state.paused = False
        
        # Build list of sections to process from approved outline
        sections_to_process = []
        for unit in st.session_state.approved_outline:
            for section in unit.get('sections', []):
                sections_to_process.append({
                    'unit_number': unit['unit_number'],
                    'unit_title': unit['unit_title'],
                    'section_number': section['section_number'],
                    'section_title': section['section_title'],
                    'description': section.get('description', '')
                })
        st.session_state.sections_to_process = sections_to_process
        
    
    total = len(st.session_state.sections_to_process)
    completed = len(st.session_state.content) # Number of successfully generated sections
    
    # Progress metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("✅ Completed", f"{completed}/{total}")
    with col2:
        progress_pct = (completed / total * 100) if total > 0 else 0
        st.metric("📊 Progress", f"{progress_pct:.0f}%")
    with col3:
        st.metric("⏳ Remaining", total - completed)
    with col4:
        if completed > 0 and not st.session_state.paused:
            elapsed = time.time() - st.session_state.generation_start_time
            avg_time = elapsed / completed
            eta_seconds = int(avg_time * (total - completed))
            eta_minutes = eta_seconds // 60
            st.metric("⏱️ ETA", f"~{eta_minutes}min")
        else:
            st.metric("⏱️ ETA", "N/A")
    
    # Progress bar
    st.progress(completed / total if total > 0 else 0)
    
    st.divider()
    
    # Generate content section by section
    if completed < total:
        # Find the next section to process that hasn't been fully generated
        current_section_data = None
        for section_data in st.session_state.sections_to_process:
            sec_key = f"{section_data['section_number']} {section_data['section_title']}"
            if sec_key not in st.session_state.content:
                current_section_data = section_data
                break
        
        if not current_section_data:
            st.error("Internal Error: Could not find next section to process.")
            return

        section_key = f"{current_section_data['section_number']} {current_section_data['section_title']}"
        
        st.subheader(f"🤖 Generating: {section_key}")
        
        col1, col2 = st.columns([3, 1])
        with col1:
            st.info(f"**Unit {current_section_data['unit_number']}:** {current_section_data['unit_title']}")
            st.write(f"**Description:** {current_section_data['description']}")
        
        with col2:
            if st.button("⏸️ Pause Generation", key="pause_gen", use_container_width=True):
                st.session_state.paused = True
                st.rerun()
        
        # Generate content if not paused
        if not st.session_state.paused:
            with st.spinner(f"✍️ Writing content for section {completed + 1} of {total}... This may take 30-60 seconds"):
                
                # Build context
                context = {
                    'course_title': st.session_state.course_title,
                    'course_code': st.session_state.course_code,
                    'credits': st.session_state.credits,
                    'target_audience': st.session_state.target_audience,
                    'program_objectives': st.session_state.program_objectives,
                    'program_outcomes': st.session_state.program_outcomes,
                    'course_outcomes': st.session_state.course_outcomes,
                    'specialized_outcomes': st.session_state.specialized_outcomes,
                    'learning_objectives': st.session_state.learning_objectives,
                    'prerequisites': st.session_state.prerequisites,
                    'assessment_methods': st.session_state.assessment_methods
                }
                
                # Generate content
                with st.expander("🔍 Content Generation Details", expanded=True):
                    content = generate_content(current_section_data, context)
                
                if content and len(content.strip()) > 100:
                    st.session_state.content[section_key] = content
                    st.session_state.content_status[section_key] = "Generated"
                    st.success(f"✅ Content generated for {section_key}")
                    
                    # Show preview
                    with st.expander("📄 Content Preview", expanded=False):
                        st.write(content[:500] + "...")
                    
                    st.divider()
                    
                    # ===== IMAGE SECTION =====
                    st.subheader("🖼️ Add Image for This Section (Optional)")
                    
                    tab1, tab2 = st.tabs(["📤 Upload Image", "🤖 Generate Image Prompt"])
                    
                    with tab1:
                        st.info("💡 Upload a relevant image for this section")
                        # Use a unique key for each section's uploader
                        uploaded_image = st.file_uploader(
                            f"Upload image for {section_key}",
                            type=['png', 'jpg', 'jpeg'],
                            key=f"image_upload_{section_key}" # Keyed by section_key
                        )
                        
                        if uploaded_image:
                            st.session_state.images[section_key] = uploaded_image
                            st.success("✅ Image uploaded successfully!")
                            st.image(uploaded_image, caption=section_key, width=300)
                    
                    with tab2:
                        st.info("💡 Generate an AI prompt to create an image for this section")
                        
                        # Initialize image prompt if not present
                        if section_key not in st.session_state.image_prompts:
                            st.session_state.image_prompts[section_key] = ""

                        if st.button("🤖 Generate Image Prompt", key=f"gen_img_prompt_{section_key}"):
                            with st.spinner("Generating image prompt..."):
                                img_prompt = generate_image_prompt_for_section(current_section_data, context)
                                if img_prompt:
                                    st.session_state.image_prompts[section_key] = img_prompt
                                    st.rerun()
                        
                        # Display and allow editing of the prompt
                        if st.session_state.image_prompts.get(section_key):
                            st.success("✅ Image prompt generated!")
                            prompt_text = st.text_area(
                                "Image Generation Prompt (edit if needed):",
                                value=st.session_state.image_prompts[section_key],
                                height=150,
                                key=f"img_prompt_text_{section_key}" # Keyed by section_key
                            )
                            # Update session state if prompt text is changed
                            if prompt_text != st.session_state.image_prompts[section_key]:
                                st.session_state.image_prompts[section_key] = prompt_text
                            
                            st.info("💡 Copy this prompt and use it with DALL-E, Midjourney, or Stable Diffusion to generate an image, then upload it in the 'Upload Image' tab")
                            
                            if st.button("📋 Copy Prompt", key=f"copy_prompt_{section_key}"):
                                st.code(prompt_text)
                    
                    st.divider()
                    
                    # Continue or skip buttons
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("⏭️ Skip Image & Continue", type="primary", use_container_width=True, key=f"skip_img_{section_key}"):
                            time.sleep(0.5)
                            st.rerun()
                    
                    with col2:
                        if st.button("✅ Save & Continue", type="primary", use_container_width=True, key=f"continue_{section_key}"):
                            time.sleep(0.5)
                            st.rerun()
                    
                else:
                    st.error("❌ Content generation failed or returned insufficient content")
                    st.session_state.content_status[section_key] = "Failed"
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Retry", type="primary", use_container_width=True, key=f"retry_{section_key}"):
                            st.rerun()
                    with col2:
                        if st.button("⏭️ Skip This Section", use_container_width=True, key=f"skip_section_{section_key}"):
                            st.session_state.content[section_key] = "[Content generation skipped]"
                            st.session_state.content_status[section_key] = "Skipped"
                            st.rerun()
        else:
            # Paused state
            st.warning("⏸️ Content generation paused")
            st.info(f"Currently at section {completed + 1} of {total}")
            
            if st.button("▶️ Resume Generation", type="primary", key="resume_gen"):
                st.session_state.paused = False
                st.rerun()
    
    else:
        # All content generated
        st.success("🎉 All Content Generated Successfully!")
        
        total_words = sum(len(c.split()) for c in st.session_state.content.values())
        est_pages = total_words // 300 + 1 if total_words > 0 else 0
        images_added = len(st.session_state.images)
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("✅ Completed", f"{completed}/{total}")
        with col2:
            st.metric("📊 Progress", "100%")
        with col3:
            st.metric("⏳ Remaining", 0)
        with col4:
            st.metric("⏱️ ETA", "Done")
            
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📚 Units", len(st.session_state.approved_outline))
        with col2:
            st.metric("📝 Total Sections", total)
        with col3:
            st.metric("📄 Est. Pages", f"~{est_pages}")
        with col4:
            st.metric("🖼️ Images", images_added)
        
        # Show content summary
        with st.expander("📊 Content Summary", expanded=False):
            for unit in st.session_state.approved_outline:
                st.write(f"**Unit {unit['unit_number']}: {unit['unit_title']}**")
                for section in unit.get('sections', []):
                    sec_key = f"{section['section_number']} {section['section_title']}"
                    content_words = len(st.session_state.content.get(sec_key, '').split())
                    status = st.session_state.content_status.get(sec_key, "Unknown")
                    has_image = "🖼️" if sec_key in st.session_state.images else ""
                    st.write(f"  - {sec_key}: {status} ({content_words:,} words) {has_image}")
        
        st.divider()
        
        # Navigation buttons
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("← Back to Outline", use_container_width=True, key="back_from_gen"):
                st.session_state.step = 'outline_generation'
                st.rerun()
        
        with col2:
            if st.button("🔄 Regenerate All", use_container_width=True, key="regen_all"):
                if st.checkbox("⚠️ Confirm: This will delete all generated content", key="confirm_regen"):
                    st.session_state.content = {}
                    st.session_state.content_status = {}
                    st.session_state.images = {}
                    st.session_state.image_prompts = {}
                    st.session_state.generation_start_time = time.time() # Reset timer
                    st.rerun()
        
        with col3:
            if st.button("📄 Compile Documents →", type="primary", use_container_width=True, key="go_compile"):
                st.session_state.step = 'compilation'
                st.rerun()

# ============================================================================
# PAGE 5: COMPILATION AND DOWNLOAD
# ============================================================================

def show_compilation_page():
    """Compilation page with PDF/DOCX generation and direct download"""
    st.header("📄 Step 5: Compile & Download")
    
    # Validation
    if 'content' not in st.session_state or not st.session_state.content:
        st.error("❌ No content to compile")
        if st.button("← Back to Content Generation", key="back_no_content"):
            st.session_state.step = 'content_generation'
            st.rerun()
        return
    
    if 'approved_outline' not in st.session_state:
        st.error("❌ No outline found")
        if st.button("← Back to Outline", key="back_no_outline_comp"):
            st.session_state.step = 'outline_generation'
            st.rerun()
        return
    
    # ========== Summary ==========
    st.subheader("📊 Content Summary")
    
    total_sections = len(st.session_state.content)
    total_words = sum(len(c.split()) for c in st.session_state.content.values())
    est_pages = total_words // 300 + 1 if total_words > 0 else 0
    images_count = len(st.session_state.images)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📚 Units", len(st.session_state.approved_outline))
    with col2:
        st.metric("📝 Sections", total_sections)
    with col3:
        st.metric("📄 Words", f"{total_words:,}")
    with col4:
        st.metric("🖼️ Images", images_count)
    
    st.info(f"📄 Estimated pages: ~{est_pages}")
    
    st.divider()
    
    # ========== Compilation Options ==========
    st.subheader("⚙️ Compilation Options")
    
    col1, col2 = st.columns(2)
    with col1:
        compile_type = st.radio(
            "What to compile:",
            ["Separate Unit Files", "Complete Course File", "Both (Separate + Complete)"],
            index=2, # Default to "Both"
            key="compile_type",
            help="Choose whether to compile each unit separately, one complete file, or both"
        )
        st.session_state.compile_type = compile_type # Update session state
    
    with col2:
        output_format = st.radio(
            "Output Format:",
            ["PDF", "DOCX (Editable)"],
            index=0, # Default to PDF
            key="output_format",
            help="PDF for final documents, DOCX for editable Word files"
        )
        st.session_state.output_format = output_format # Update session state
    
    
    st.divider()
    
    # ========== Compilation Button ==========
    if st.button("🔨 Start Compilation", type="primary", use_container_width=True, key="start_compile"):
        
        course_info = {
            'course_title': st.session_state.course_title,
            'course_code': st.session_state.course_code,
            'credits': st.session_state.credits,
            'target_audience': st.session_state.target_audience
        }
        
        
        compiled_files = {}
        
        # ===== Compile Separate Units =====
        if compile_type in ["Separate Unit Files", "Both (Separate + Complete)"]:
            st.subheader("📚 Compiling Individual Units")
            progress_bar_units = st.progress(0)
            status_text_units = st.empty()
            
            for i, unit in enumerate(st.session_state.approved_outline):
                unit_num = unit['unit_number']
                # Ensure unit title is present and clean it for filename
                unit_title = unit.get('unit_title', f"Unit_{unit_num}")
                safe_unit_title = re.sub(r'[^\w\s-]', '', unit_title).strip()
                if not safe_unit_title: safe_unit_title = f"Unit_{unit_num}"
                
                status_text_units.text(f"Compiling Unit {unit_num}: {safe_unit_title}...")
                
                with st.spinner(f"Compiling Unit {unit_num}: {safe_unit_title}..."):
                    if output_format == "PDF":
                        file_buffer = compile_unit_pdf(unit, course_info, st.session_state.content)
                    else:
                        file_buffer = compile_unit_docx(unit, course_info, st.session_state.content)
                    
                    if file_buffer:
                        ext = '.pdf' if output_format == "PDF" else '.docx'
                        filename = f"Unit_{unit_num}_{safe_unit_title}{ext}"
                        mime_type = 'application/pdf' if output_format == "PDF" else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        
                        compiled_files[f"unit_{unit_num}"] = {
                            'buffer': file_buffer,
                            'filename': filename,
                            'mime_type': mime_type
                        }
                        
                        
                        st.success(f"✅ Unit {unit_num} compiled successfully")
                    else:
                        st.error(f"❌ Failed to compile Unit {unit_num}")
                
                progress_bar_units.progress((i + 1) / len(st.session_state.approved_outline))
            
            status_text_units.text("✅ All individual units compiled!")
        
        # ===== Compile Complete File =====
        if compile_type in ["Complete Course File", "Both (Separate + Complete)"]:
            st.subheader("📖 Compiling Complete Course")
            
            with st.spinner("Compiling complete course file..."):
                if output_format == "PDF":
                    file_buffer = compile_complete_pdf(
                        st.session_state.approved_outline,
                        course_info,
                        st.session_state.content
                    )
                else:
                    file_buffer = compile_complete_docx(
                        st.session_state.approved_outline,
                        course_info,
                        st.session_state.content
                    )
                
                if file_buffer:
                    ext = '.pdf' if output_format == "PDF" else '.docx'
                    # Clean course title for filename
                    safe_course_title = re.sub(r'[^\w\s-]', '', st.session_state.course_title).strip()
                    if not safe_course_title: safe_course_title = "Course"
                    filename = f"Complete_{st.session_state.course_code}_{safe_course_title}{ext}"
                    mime_type = 'application/pdf' if output_format == "PDF" else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    
                    compiled_files['complete'] = {
                        'buffer': file_buffer,
                        'filename': filename,
                        'mime_type': mime_type
                    }
                    
                    
                    st.success("✅ Complete course compiled successfully")
                else:
                    st.error("❌ Failed to compile complete course")
        
        # ===== Download Section =====
        st.divider()
        st.subheader("📥 Download Files")
        
        if compiled_files:
            st.success(f"✅ {len(compiled_files)} file(s) ready for download")
            
            # Display downloads using columns for better layout
            num_files = len(compiled_files)
            cols = st.columns(min(3, num_files)) # Max 3 columns for downloads
            col_index = 0

            # Sort files for consistent display order: complete first, then units numerically
            sorted_keys = sorted(compiled_files.keys(), key=lambda x: (0 if x == 'complete' else 1, int(x.split('_')[1]) if x.startswith('unit_') else 0))
            
            for key in sorted_keys:
                file_data = compiled_files[key]
                with cols[col_index % len(cols)]:
                    st.write("---") # Separator for each download button block
                    st.write(f"📄 {file_data['filename']}")
                    file_size = len(file_data['buffer'].getvalue()) / 1024  # KB
                    st.caption(f"Size: {file_size:.1f} KB")
                    
                    file_data['buffer'].seek(0) # Ensure buffer is at the beginning for download
                    st.download_button(
                        label="⬇️ Download",
                        data=file_data['buffer'].getvalue(),
                        file_name=file_data['filename'],
                        mime=file_data['mime_type'],
                        key=f"download_{key}_{output_format}", # Unique key
                        use_container_width=True
                    )
                col_index += 1
        else:
            st.warning("⚠️ No files were compiled")
        
        st.divider()
        
        # Navigation
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("← Back to Content", key="back_to_content_comp"):
                st.session_state.step = 'content_generation'
                st.rerun()
        with col2:
            if st.button("🔄 Compile Again", key="compile_again"):
                st.rerun()
        with col3:
            if st.button("🏠 New Project", key="new_proj"):
                # Save important data that should persist if the user starts a new project
                persistent_data = {
                    'api_key': st.session_state.api_key,
                }
                
                # Clear everything
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                
                # Reinitialize and restore persistent data
                initialize_session_state()
                for key, value in persistent_data.items():
                    st.session_state[key] = value
                
                st.success("✅ Ready for new project!")
                time.sleep(1)
                st.rerun()


# ============================================================================
# PHASE 4 PART 2 COMPLETE
# ============================================================================
print("Phase 4 Part 2: Content Generation, Compilation, and Sidebar loaded successfully")


# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    """Main application entry point"""
    
    # Page configuration
    st.set_page_config(
        page_title="AI Curriculum Generator",
        page_icon="🎓",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Title and description
    st.title("🎓 AI Curriculum Generator")
    st.caption("Professional academic materials with eGyankosh standards | Powered by Grok AI")
    
    # Initialize session state
    initialize_session_state()
    
    # Show navigation
    show_navigation()
    
    # Route to appropriate page based on current step
    step = st.session_state.step
    
    if step == 'syllabus_upload':
        show_syllabus_upload_page()
    
    elif step == 'configuration':
        show_configuration_page()
    
    elif step == 'outline_generation':
        show_outline_page()
    
    elif step == 'content_generation':
        show_content_generation_page()
    
    elif step == 'compilation':
        show_compilation_page()
    
    else:
        st.error(f"❌ Unknown step: {step}")
        st.session_state.step = 'configuration' # Default to configuration
        st.rerun()
    
    # Show sidebar status
    show_sidebar_status()
    
    # Footer
    st.divider()
    with st.expander("ℹ️ About This Application", expanded=False):
        st.markdown("""
        ### AI Curriculum Generator
        
        **Features:**
        - PDF syllabus extraction
        - AI-powered content generation
        - Image upload and prompt generation
        - Professional PDF/DOCX compilation
        - Fixed LaTeX equation rendering
        
        **Technologies:**
        - Grok AI (grok-2-1212)
        - ReportLab (PDF generation)
        - python-docx (DOCX generation)
        - PyPDF2 (PDF text extraction)
        - PIL (Image processing)
        
        **Content Generation:**
        - Learning objectives and outcomes
        - Detailed explanations with examples
        - Practice problems and solutions
        - Case studies and real-world applications
        - Assessment questions
        - LaTeX equation support (using $$...$$ format)
        
        **Output Formats:**
        - PDF (professional, print-ready)
        - DOCX (editable Word documents)
        - Separate files per unit or complete course file
        - Direct browser download
        
        **Quality Features:**
        - Comprehensive error handling
        - Progress tracking
        - Content validation
        - Image integration
        - Professional formatting
        - Academic structure
        
        **Usage:**
        1. Configure course details (title, code, audience, etc.)
        2. Optional: Upload reference PDF for context
        3. Generate and approve course outline
        4. Generate content for each section
        5. Generate image prompts for visual content
        6. Compile and download final documents
        
        **Requirements:**
        pip install streamlit requests pillow reportlab python-docx pypdf2
        
        Note: Make sure all required packages are installed for full functionality.
        """)

# ============================================================================
# APPLICATION ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    main()

# ============================================================================
# INSTALLATION REQUIREMENTS
# ============================================================================
"""
Required packages - install with pip:

pip install streamlit
pip install requests
pip install pillow
pip install PyPDF2
pip install reportlab
pip install python-docx

Or use requirements.txt:

streamlit>=1.29.0
requests>=2.31.0
pillow>=10.0.0
PyPDF2>=3.0.0
reportlab>=4.0.0
python-docx>=1.1.0
"""

# ============================================================================
# HOW TO RUN
# ============================================================================
"""
1. Save all phases to a single file named 'curriculum_generator.py':
   - Copy Phase 1 content
   - Copy Phase 2 content
   - Copy Phase 3 content
   - Copy Phase 4 Part 1 content
   - Copy Phase 4 Part 2 content
   - Copy This file: Main Application

2. Install requirements:
   pip install -r requirements.txt

3. Run the application:
   streamlit run curriculum_generator.py

4. Open browser at:
   http://localhost:8501
"""

# ============================================================================
# TROUBLESHOOTING
# ============================================================================
"""
Common Issues and Solutions:

1. **LaTeX equations not rendering:**
   - Fixed! Now converts LaTeX to Unicode symbols
   - Example: \\leq becomes ≤

2. **API errors:**
   - Check API key starts with 'xai-'
   - Test API using the test button
   - Check internet connection

3. **PDF compilation fails:**
   - Check ReportLab is installed
   - Try DOCX format instead
   - Check error logs in UI

4. **Images not appearing:**
   - Ensure images are PNG/JPG/JPEG
   - Check file size < 5MB
   - Upload one image at a time

5. **Content too short:**
   - Check API logs in UI
   - Increase max_tokens if needed
   - Check API rate limits
"""

# ============================================================================
# COMPLETE APPLICATION - READY TO USE
# ============================================================================
print("✅ All components loaded successfully!")
print("🎓 AI Curriculum Generator ready!")
print("🚀 Run with: streamlit run curriculum_generator.py")
